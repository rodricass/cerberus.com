"""
Definition of views.
"""

from django.shortcuts import render
from django.http import HttpRequest
from django.http import Http404
from django.template import RequestContext
from datetime import datetime
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.core.urlresolvers import reverse, reverse_lazy
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
from django.contrib.auth.models import User
import os
import posixpath
import re
import hashlib
import json
import time
from django.utils.safestring import SafeString
from docx import Document
from docx.shared import Pt
from django.db.models import Q
from django.views.generic.edit import FormView

from .models import *
from .forms import InvestigacionForm, DocumentoForm, InvestigacionFormEdit, BuscadorGeneralForm, BuscadorInvestigacionesForm, NotaForm, UsuariosForm, UsuarioNuevoForm, RegexForm
from .functions import *

from .myclasses import *
from collections import defaultdict
from .genericos import *
from .fileuploader import *

#***********************************************************************************************************************************************************************
#********************************************************************************** START INICIO ***********************************************************************
#***********************************************************************************************************************************************************************

@login_required
def home(request):
    """Crea la página de inicio"""
    assert isinstance(request, HttpRequest)
    mensajes = Mensaje.objects.filter(eliminado=False).filter(receptor=request.user).order_by('-fecha_agregado')
    investigaciones = Investigacion.objects.filter(usuario=request.user).filter(eliminado=False).filter(finalizado_incorrecto=False).filter(finalizado_correcto=False).order_by('-fecha_agregado')[:2]
    documentos = Documento.objects.filter(usuario=request.user).filter(eliminado=False).order_by('-fecha_agregado')[:7]
    form = DocumentoForm()
    form_usuario = UsuariosForm(request.user)
    context = {
            'title':'Home',
            'year':datetime.now().year,
            'mensajes':mensajes,
            'investigaciones':investigaciones,
            'formDoc':form,
            'form_usuario':form_usuario,
            'documentos':documentos,
            }
    if request.user.is_superuser:
        form_crear = UsuarioNuevoForm()
        form_usuario = UsuariosForm(request.user)
        context['form_crear'] = form_crear
        context['form_usuario'] = form_usuario
    return render(request,'FrontEnd/index.html',context)

def crear_usuario(request):
    """Crea usuario nuevo"""
    if request.method == 'POST':
        form = UsuarioNuevoForm(request.POST)
        if form.is_valid():
            nombre = form.cleaned_data['nombre']
            contraseña = form.cleaned_data['contraseña']
            user = User.objects.create_user(nombre, '', contraseña)
            user.save()

    return HttpResponseRedirect(reverse('home'))

def eliminar_usuario(request):
    """Elimina usuario existente"""
    if request.method == 'POST':
        usuario = request.POST.get('usuarios')
        usuario_instance = User.objects.get(id=usuario)
        usuario_instance.delete()

    return HttpResponseRedirect(reverse('home'))

def logout_view(request):
    """Cierra sesión de usuario"""
    logout(request)
    return HttpResponseRedirect(reverse('home'))

#***********************************************************************************************************************************************************************
#********************************************************************************** END INICIO *******************************************************************
#***********************************************************************************************************************************************************************


#***********************************************************************************************************************************************************************
#********************************************************************************** START DOCUMENTOS *******************************************************************
#***********************************************************************************************************************************************************************

@login_required
def documentos(request):
    """Genera la vista de todos los documentos subidos por el usuario"""
    form = BuscadorInvestigacionesForm(request.user)
    formDoc = DocumentoForm()

    context = {'form':form, 
               'formDoc':formDoc,
               'title':'Documentos'
               }
    if request.method == 'POST':
        id_investigacion = request.POST.get('investigaciones')
        investigacion = Investigacion.objects.get(id=id_investigacion)
        camino = f'Documentos:documentos>{investigacion.nombre}:investigacion|{investigacion.id}'
        documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
        context['id_investigacion'] = id_investigacion
        context['documentos'] = documentos
        context['nombre_investigacion'] =investigacion.nombre
        context['inicial'] = False
        context['camino'] = camino
    return render(request,'FrontEnd/documentos.html',context)

@login_required
def documentos_investigacion(request,investigacion_id,destino,camino):
    """Genera la vista de todos los documentos subidos por el usuario para una investigación """
    formDoc = DocumentoForm()
    context = {
                'formDoc':formDoc,
                'title':'Documentos de investigacion'
               }
    breadcrumb_path = parser_camino(camino)
    investigacion = Investigacion.objects.get(id=investigacion_id)
    documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
    context['id_investigacion'] = investigacion_id
    context['nombre_investigacion'] = investigacion.nombre
    context['documentos'] = documentos
    context['camino'] = camino
    context['camino_array'] = breadcrumb_path
    return render(request,destino,context)

@login_required
def agregar_doc(request, investigacion_id):
    """ Crear un documento en base a un archivo cargado en el sistema por el usuario"""
    #Esta vista solo es accedida desde una investigacion específica o desde el inicio
    if request.method == 'POST':
        #Crea el objeto documento que solo contiene la ubicación del documento en el servidor
        next = request.POST.get('next', '/')
        tipo = request.POST.get('tipo_archivo', '')
        creador = ObtenerCreador().creadorArchivo(tipo,request,investigacion_id)
        creador.cargar()     
        if next == "investigacion":
            return HttpResponseRedirect(reverse('investigacion',args=[investigacion_id]))
        else:
            return HttpResponseRedirect(reverse('home'))
    

#class FileFieldView(FormView):
#    form_class = DocumentoForm
#    success_url = ''

#    def post(self,request,*args,**kwargs):
#        form_class = self.get_form_class()
#        form = self.get_form(form_class)
#        next = request.POST.get('next', '/')
#        self.success_url = reverse_lazy(next)
#        investigacion_id = self.kwargs['investigacion_id']

#        if form.is_valid():
#            crear_documento_general(request,form,investigacion_id)
#            return self.form_valid(form)
#        else:
#            return self.form_invalid(form)

@login_required
def agregar_docDocumentos(request, investigacion_id):
    """Crear documento desde la pantalla de documentos en base a un archivo cargado en el sistema por el usuario"""
    
    #Crea el objeto documento que solo contiene la ubicación del documento en el servidor
    tipo = request.POST.get('tipo_archivo', '')
    creador = ObtenerCreador().creadorArchivo(tipo,request,investigacion_id)
    creador.cargar()
    form = BuscadorInvestigacionesForm(request.user)
    formDoc = DocumentoForm()
    context = {'form':form, 'formDoc':formDoc }
    investigacion = Investigacion.objects.get(id=investigacion_id)
    documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
    context['id_investigacion'] = investigacion_id
    context['documentos'] = documentos
    context['inicial'] = False
    return render(request,'FrontEnd/documentos.html',context)

@login_required
def agregar_docInvestigacion(request,investigacion_id,camino):
    """Crear documento desde la pantalla de documentos de una investigacion en base a un archivo cargado en el sistema por el usuario"""
    tipo = request.POST.get('tipo_archivo', '')
    creador = ObtenerCreador().creadorArchivo(tipo,request,investigacion_id)
    creador.cargar()

    destino = 'FrontEnd/documentos_investigacion.html'

    return HttpResponseRedirect(reverse('documentos_investigacion',args=[investigacion_id,destino,camino]))

def eliminar_documentos_general(id_doc):
    """Eliminar documentos"""
    doc = Documento.objects.get(id=id_doc)
    doc.eliminado = True
    doc.save()
    parrafos = Parrafo.objects.filter(doc=doc)
    for parrafo in parrafos:
        parrafo.eliminado = True
        parrafo.save()
    entidades = EntidadesDoc.objects.filter(doc=doc)
    for entidad in entidades:
        entidad.eliminado = True
        entidad.save()
    tokens = TokensDoc.objects.filter(doc=doc)
    for token in tokens:
        token.eliminado = True
        token.save()
    notas = NotaDocumento.objects.filter(entidad=doc)
    for nota in notas:
        nota.eliminado = True
        nota.save()
    mensajes = Mensaje.objects.filter(documento=doc)
    for mensaje in mensajes:
        mensaje.eliminado = True
        mensaje.save()

@login_required
def eliminar_doc(request, id_doc, id_investigacion):
    """Eliminar un documento de la investigacion"""
    next = ''
    if request.method == 'POST':
        next = request.POST.get('next', '/')
        eliminar_documentos_general(id_doc)
    investigacion = Investigacion.objects.get(id=id_investigacion)
    documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
    form = BuscadorInvestigacionesForm(request.user)
    formDoc = DocumentoForm()
    context = {'form':form,
               'formDoc':formDoc,
               'documentos':documentos,
               'inicial':False,
               'id_investigacion':id_investigacion}
    return render(request,next,context)

@login_required
def eliminar_doc_investigacion(request,id_doc,id_investigacion,camino):
    """Elimina un documento desde la pantalla que muestras los documentos de una investigacion específico"""
    
    eliminar_documentos_general(id_doc)

    destino = 'FrontEnd/documentos_investigacion.html'

    return HttpResponseRedirect(reverse('documentos_investigacion',args=[id_investigacion,destino,camino]))



def eliminar_docInicio(request,id_doc):
    """Elimina un documento desde la pantalla de inicio"""

    eliminar_documentos_general(id_doc)
    
    return HttpResponseRedirect(reverse('home'))

@login_required
def ver_doc(request, id_doc):
    """Descarga el documento para que el usuario pueda leerlo desde su dispositivo"""
    
    doc = Documento.objects.get(id=id_doc)
    nombre, punto, ext = str(doc.documento).rpartition(".")

    content = ExtensionArchivo().getResponseDoc(ext)

    response = HttpResponse(doc.documento, content_type=content)
    file_path = f'{os.path.dirname(os.path.dirname(os.path.abspath(__file__)))}/documentos{doc.documento}'
    response['Content-Disposition'] = 'attachment; filename=%s' % doc.nombre_doc

    return response

@login_required
def mensaje_nuevo(request, id_doc, id_investigacion):
    """Genera una solicitud de eliminación para el propietario del documento"""
    form = BuscadorInvestigacionesForm(request.user)
    formDoc = DocumentoForm()
    context = {'form':form,
               'formDoc': formDoc,}
    if request.method == 'POST':
        next = request.POST.get('next', '/')
        msj = request.POST.get('mensaje')
        doc = Documento.objects.get(id = id_doc)
        investigacion = Investigacion.objects.get(id=id_investigacion)
        mensaje = Mensaje()
        mensaje.emisor = request.user
        mensaje.receptor = doc.propietario_doc
        mensaje.mensaje = msj
        mensaje.documento = doc
        mensaje.eliminado = False
        mensaje.investigacion = investigacion
        mensaje.save()
        documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
        context['documentos'] = documentos
        context['inicial'] = False
        context['id_investigacion'] = id_investigacion
    return render(request,next,context)

@login_required
def eliminar_mensaje(request,id_msj):
    """Elimina un mensaje de solicitud de eliminación de documento"""
    mensaje = Mensaje.objects.get(id=id_msj)
    mensaje.eliminado = True
    mensaje.save()
    
    return HttpResponseRedirect(reverse('home'))



#***********************************************************************************************************************************************************************
#********************************************************************************** END DOCUMENTOS *******************************************************************
#***********************************************************************************************************************************************************************

#***********************************************************************************************************************************************************************
#****************************************************************************************** START INVESTIGACIONES ****************************************************************
#***********************************************************************************************************************************************************************

@login_required
def investigaciones(request):
    """Genera la vista de las investigaciones de un usuario determinado"""
    
    investigaciones = Investigacion.objects.filter(usuario=request.user).filter(eliminado=False).filter(finalizado_incorrecto=False).filter(finalizado_correcto=False).order_by('-fecha_agregado')
    context = {'investigaciones':investigaciones, 
               'title':'Investigaciones en curso',
               }
    return render(request,'FrontEnd/investigaciones.html',context)

@login_required
def investigaciones_finalizadas(request):
    """Genera la vista de los investigaciones finalizadas de un usuario determinado"""
    
    investigaciones = Investigacion.objects.filter(usuario=request.user).filter(eliminado=False).filter(Q(finalizado_correcto=True) | Q(finalizado_incorrecto=True)).order_by('-fecha_agregado')
    form_usuario = UsuariosForm(request.user)
    context = {'investigaciones':investigaciones, 
               'form_usuario':form_usuario,
               'destino_resultados':'FrontEnd/resultados_investigacionfinalizada.html',
               'destino_informes':'FrontEnd/informes_investigacionfinalizada.html',
               'destino_documentos':'FrontEnd/documentos_investigacionfinalizada.html',
               'title':'Investigaciones finalizadas',
               }
    return render(request,'FrontEnd/investigaciones_finalizadas.html',context)

@login_required
def investigacion_finalizada(request,investigacion_id):
    """Genera la vista que muestra toda la información específica de una investigacion finalizada"""
    comprobar_usuario(request.user,investigacion_id)
    investigacion = Investigacion.objects.get(id=investigacion_id)
    form_usuario = UsuariosForm(request.user)
    camino = f'Investigaciones finalizadas:investigaciones_finalizadas>{investigacion.nombre}:investigacion_finalizada|{investigacion.id}'

    context = {'investigacion':investigacion, 
               'form_usuario':form_usuario,
               'destino_resultados':'FrontEnd/resultados_investigacionfinalizada.html',
               'destino_informes':'FrontEnd/informes_investigacionfinalizada.html',
               'destino_documentos':'FrontEnd/documentos_investigacionfinalizada.html',
               'title':f'{investigacion.nombre[:12]}...',
               'camino':camino,
               'modelo': TipoModelo().getModeloString(investigacion.modelo),
               }
    return render(request,'FrontEnd/investigacion_finalizada.html',context)

@login_required
def investigacion(request,investigacion_id):
    """Genera la vista que muestra toda la información específica a un investigacion"""
    comprobar_usuario(request.user,investigacion_id)
    investigacion = Investigacion.objects.get(id=investigacion_id)
    form = DocumentoForm()
    form_usuario = UsuariosForm(request.user)
    camino = f'Investigaciones en curso:investigaciones>{investigacion.nombre}:investigacion|{investigacion.id}'

    context = {'formDoc':form, 
               'form_usuario':form_usuario,
               'destino_resultados':'FrontEnd/resultados_investigacion.html',
               'destino_informes':'FrontEnd/informes_investigacion.html',
               'destino_documentos':'FrontEnd/documentos_investigacion.html',
               'camino':camino,
               'investigacion': investigacion,
               'title': f'{investigacion.nombre[:12]}...',
               'modelo': TipoModelo().getModeloString(investigacion.modelo),
              }
    return render(request,'FrontEnd/investigacion.html',context)


@login_required
def editar_investigacion(request,investigacion_id,camino):
    """ Editar una investigación ya existente"""
    investigacion = Investigacion.objects.get(id=investigacion_id)
    if request.method == 'POST':
        #Actualizar el investigacion
        form = InvestigacionFormEdit(instance=investigacion, data=request.POST)
        if form.is_valid():
            form.save()
            return HttpResponseRedirect(reverse('investigaciones'))
    else:
        #Solicita el investigacion para ser modificado, con información precargada
        form = InvestigacionFormEdit(instance=investigacion)
    
    breadcrumb_path = parser_camino(camino)
    context = {'form':form, 
               'investigacion': investigacion, 
               'investigacion_nombre':investigacion.nombre, 
               'documentos':documentos,
               'title':'Editar investigacion',
               'camino':camino,
               'camino_array':breadcrumb_path,
               }

    return render(request,'FrontEnd/editar_investigacion.html',context)


def eliminar_generico_investigaciones(request,investigacion):
    """Función para eliminar investigaciones"""
    if investigacion.propietario == request.user:
        documentos = Documento.objects.filter(investigacion=investigacion)
        for doc in documentos:
            parrafos = Parrafo.objects.filter(doc=doc)
            for parrafo in parrafos:
                parrafo.eliminado = True
                parrafo.save()
            entidades = EntidadesDoc.objects.filter(doc=doc)
            for entidad in entidades:
                entidad.eliminado = True
                entidad.save()
            tokens = TokensDoc.objects.filter(doc=doc)
            for token in tokens:
                token.eliminado = True
                token.save()
            notas = NotaDocumento.objects.filter(entidad=doc)
            for nota in notas:
                nota.eliminado = True
                nota.save()
            mensajes = Mensaje.objects.filter(documento=doc)
            for mensaje in mensajes:
                mensaje.eliminado = True
                mensaje.save()
            doc.eliminado = True
            doc.save()
        investigacion.eliminado = True
        investigacion.save()
    else:
        investigacion.usuario.remove(request.user.id)

@login_required
def eliminar_finalizada(request,investigacion_id):
    """Elimina una investigacion finalizada específico"""

    investigacion = Investigacion.objects.get(id=investigacion_id)
    eliminar_generico_investigaciones(request,investigacion)

    return HttpResponseRedirect(reverse('investigaciones_finalizadas'))

@login_required
def administrar_investigaciones(request,tipo):
    """Administrar investigaciones seleccionadas, para ser eliminadas o finalizadas"""
    if request.method == 'POST':
        ids = request.POST.getlist('checks[]')
        if tipo == "incorrecto":
            #Se han seleccionado investigaciones para ser finalizadas incorrectamente
            for id in ids:
                 investigacion = Investigacion.objects.filter(id=id)
                 for c in investigacion:
                    if c.propietario == request.user:
                        c.finalizado_incorrecto = True
                        c.save()
                    else:
                        c.usuario.remove(request.user.id)
        else:
            #Se han seleccionado investigaciones para ser finalizadas correctamente
            for id in ids:
                investigacion = Investigacion.objects.filter(id=id)
                for c in investigacion:
                    if c.propietario == request.user:
                        c.finalizado_correcto = True
                        c.save()
                    else:
                        c.usuario.remove(request.user.id)
    
    return HttpResponseRedirect(reverse('investigaciones'))


@login_required
def nueva_investigacion(request,camino):
    """Crea una nueva investigación"""
    if request.method != 'POST':
        #No se ha enviado información, se crea una forma vacía
        form = InvestigacionForm()
    else:
        #POST información recibida, se procesa la misma
        form = InvestigacionForm(request.POST)
        if form.is_valid():
            nueva_investigacion = form.save(commit=False)
            nueva_investigacion.propietario = request.user
            nueva_investigacion.eliminado = False
            nueva_investigacion.finalizado_correcto = False
            nueva_investigacion.finalizado_incorrecto = False
            nueva_investigacion.save()
            nueva_investigacion.usuario.add(request.user)
            return HttpResponseRedirect(reverse('investigaciones'))
    
    breadcrumb_path = parser_camino(camino)
    context = {'form':form,
               'title':'Nueva investigación',
               'camino':camino,
               'camino_array':breadcrumb_path,
               }
    return render(request,'FrontEnd/nueva_investigacion.html',context)

@login_required
def compartir_investigaciones(request,investigacion_id,tipo):
    """Comparte un investigacion a otro usuario"""
    if request.method == 'POST':
        next = request.POST.get('next', '/')
        usuario = request.POST.get('usuarios')
        usuario_instance = User.objects.get(id=usuario)
        investigacion = Investigacion.objects.get(id=investigacion_id)
        if not(investigacion.usuario.filter(id=usuario_instance.id)):
            investigacion.usuario.add(usuario)
        if tipo=="con":
            if investigacion.propietario != usuario_instance:
                investigacion.propietario = usuario_instance
                investigacion.save()
    if next == "investigacion":
        return HttpResponseRedirect(reverse(next,args=[investigacion_id]))
    else:
        return HttpResponseRedirect(reverse(next))

@login_required
def compartir_investigacionFinalizada(request,investigacion_id):
    """Comparte una investigación finalizada a otro usuario"""
    if request.method == 'POST':
        usuario = request.POST.get('usuarios')
        usuario_instance = User.objects.get(id=usuario)
        investigacion = Investigacion.objects.get(id=investigacion_id)
        if not(investigacion.usuario.filter(id=usuario_instance.id)):
            investigacion.usuario.add(usuario)

    return HttpResponseRedirect(reverse('investigacion_finalizada'),args=[investigacion_id])  
    

#***********************************************************************************************************************************************************************
#****************************************************************************************** END INVESTIGACIONES ******************************************************************
#***********************************************************************************************************************************************************************

#***********************************************************************************************************************************************************************
#****************************************************************************************** START NOTAS ****************************************************************
#***********************************************************************************************************************************************************************

def comprobar_usuario(usuario,investigacion_id):
    """Comprueba que el usuario que solicita la investigacion sea usuario del mismo"""
    investigacion = Investigacion.objects.get(id=investigacion_id)
    if not(investigacion.usuario.filter(id=usuario.id)):
        raise Http404("Permiso denegado")

@login_required
def notas(request, id, tipo, camino):
    """Genera la vista de todas las notas"""
    notas = {
             "documento": NotaDocumento.objects.filter(entidad=id).filter(eliminado=False).order_by('-fecha_agregado'),
             "investigacion": NotaInvestigacion.objects.filter(entidad=id).filter(eliminado=False).order_by('-fecha_agregado'),
             }
    context = { "notas" : notas[tipo],
                "id": id,
                "tipo": tipo,
                "title":'Notas',
                }
    if request.method != 'POST':
        form = NotaForm()
        breadcrumb_path = parser_camino(camino)
    else:
        crearNota = {
                    "documento": NotaDocumento,
                    "investigacion": NotaInvestigacion,
            }
        form = NotaForm(request.POST)
        nueva_nota = form.save(commit=False)
        notaHija = crearNota[tipo](nota=nueva_nota.nota, descripcion=nueva_nota.descripcion, eliminado=False)
        if tipo == "documento":
            doc = Documento.objects.get(id=id)
            notaHija.entidad = doc
        elif tipo == "investigacion":
            notaHija.entidad = Investigacion.objects.get(id=id)
        notaHija.save()
        return HttpResponseRedirect(reverse('notas',args=[id,tipo,camino]))

    context['form'] = form
    context['camino'] = camino
    context['camino_array'] = breadcrumb_path
    
    return render(request,'FrontEnd/notas.html',context)
    
def parser_camino(camino):
    #Un camino es enviado como medio para reutilizar una pantalla cambiando la breadcrumbbs nav en base a desde donde es accedida
    #El > separa distintas referencias a páginas para poner en el breadcrumbb nav. Dentro de una referencia el : separa el nombre de la pagina de la parte correspondiente
    # a la url necesaria para dirigirse a la misma. La parte derecha del : corresponde al nombre correspondiente a la url y se encuentra separado de los parametros por
    # una barra |. Los parametros se encuentran uno al lado del otro separados por comas
    camino_aux = camino.split(">")
    breadcrumb_path = []
    for segmento in camino_aux:
        if segmento[0] == '*':
            breadcrumb_path.append([segmento[1:],True])
        elif segmento[0] != " ":
            segmento_aux = segmento.split(":")
            url = segmento_aux[1].split("|")
            if len(url) == 1:
                parametros = " "
            else:
                parametros = url[1]
            breadcrumb_path.append([segmento_aux[0],False,url[0],parametros])
    return breadcrumb_path

@login_required
def eliminar_nota(request,id,tipo,id_nota,camino):
    """Elimina una nota desde la vista específica que muestra las notas de las investigaciones o documentos"""
    if tipo == "documento":
        nota_eliminar = NotaDocumento.objects.get(id=id_nota)
    elif tipo == "investigacion":
        nota_eliminar = NotaInvestigacion.objects.get(id=id_nota)
    nota_eliminar.eliminado = True
    nota_eliminar.save()
    return HttpResponseRedirect(reverse('notas',args=[id,tipo,camino]))

@login_required
def eliminar_notainvestigacion(request,id_investigacion,id_nota):
    """Elimina una nota de investigación y retorna a la vista asociada a la sección de notas genérica"""
    form_elegir = BuscadorInvestigacionesForm(request.user)
    form_crear = NotaForm()
    investigacion = Investigacion.objects.get(id=id_investigacion)
    documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
    notas = NotaInvestigacion.objects.filter(entidad=id_investigacion).filter(eliminado=False).order_by('-fecha_agregado')
    context = {
                'form_elegir':form_elegir,
                'form_crear': form_crear,
                'id_investigacion' : id_investigacion,
                'documentos' : documentos,
                'nombre_investigacion':investigacion.nombre,
                'inicial' : False,
                'notas' : notas,
                'title':'Notas',
               }
    nota_eliminar = NotaInvestigacion.objects.get(id=id_nota)
    nota_eliminar.eliminado = True
    nota_eliminar.save()
    return render(request,'FrontEnd/ver_notas.html',context)

@login_required
def ver_notas(request):
    """Ver notas de investigaciones o documentos"""
    form_elegir = BuscadorInvestigacionesForm(request.user)
    form_crear = NotaForm()
    context = {
                'form_elegir':form_elegir,
                'form_crear': form_crear,
                "title":'Notas de investigaciones',
               }
    if request.method == 'POST':
        id_investigacion = request.POST.get('investigaciones')
        investigacion = Investigacion.objects.get(id=id_investigacion)
        documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
        notas = NotaInvestigacion.objects.filter(entidad=id_investigacion).filter(eliminado=False).order_by('-fecha_agregado')
        context['id_investigacion'] = id_investigacion
        context['nombre_investigacion'] = investigacion.nombre
        context['documentos'] = documentos
        context['inicial'] = False
        context['notas'] = notas
    return render(request,'FrontEnd/ver_notas.html',context)
 
@login_required
def crear_nota(request, id_investigacion):
    """Crea una nota a una investigación en la pantalla de ver_notas"""
    investigacion = Investigacio.objects.get(id=id_investigacion)
    documentos = investigacion.documento_set.filter(eliminado=False).order_by('-fecha_agregado')
    notas = NotaInvestigacio.objects.filter(entidad=id_investigacion).filter(eliminado=False).order_by('-fecha_agregado')
    form_elegir = BuscadorInvestigacionesForm(request.user)
    form_crear = NotaForm()
    context = {
            'form_elegir':form_elegir,
            'form_crear': form_crear,
            'id_investigacion': id_investigacion,
            'documentos': documentos,
            'nombre_investigacion': investigacion.nombre,
            'notas': notas,
            'inicial': False,
        }
    if request.method == 'POST':
        form = NotaForm(request.POST)
        nueva_nota = form.save(commit=False)
        notaHija = NotaInvestigacio(nota=nueva_nota.nota, descripcion=nueva_nota.descripcion, eliminado=False)
        notaHija.entidad = investigacion
        notaHija.save()
    return render(request,'FrontEnd/ver_notas.html',context)

#*********************************************************************************************************************************************************************
#****************************************************************************************** END NOTAS ****************************************************************
#*********************************************************************************************************************************************************************

#*********************************************************************************************************************************************************************
#**************************************************************************************START BÚSQUEDAS ***************************************************************
#*********************************************************************************************************************************************************************

@login_required
def buscador_general(request,investigacion_id,camino):
    """Genera la vista para las búsquedas generales"""

    investigacion = Investigacion.objects.get(id=investigacion_id)
    context = {"investigacion":investigacion,
               'title':'Buscador general',
               }
    if request.method != "POST":
        #Solicita la página de búsqueda
        form = BuscadorGeneralForm()    
    else:
        #Busca en los documentos y retorna las apariciones
        
        form = BuscadorGeneralForm(request.POST)
        
        if form.is_valid():
            string = form.cleaned_data['busqueda']
            context["expresion"] = string
            documentos = investigacion.documento_set.all().order_by('-fecha_agregado')
            resultados = []
            id_fila = 0
            for documento in documentos:
                parrafos = Parrafo.objects.filter(doc=documento)
                for parrafo in parrafos:
                    aux = posiciones(parrafo.parrafo,string)
                    if len(aux) > 0:
                        resultados.append([documento.nombre_doc,documento.id,parrafo.nro,parrafo.parrafo,aux,id_fila])
                        id_fila += 1
            context["res"] = resultados
            resultados_json = json.dumps(resultados)
            context["json"] = resultados_json
    context["form"] = form
    breadcumb_path = parser_camino(camino)
    context["camino"] = camino
    context["camino_array"] = breadcumb_path
    return render(request,'FrontEnd/buscador_general.html',context)

def posiciones(parrafo,palabra):
    """Obtiene las posiciones en las cuáles una palabra determinada aparece en un párrafo"""

    respuesta = []
    posicion_actual = 0
    palabra = palabra.lower()
    parrafo = parrafo.lower()
    aparicion = parrafo.find(palabra)
    long = len(palabra)
    while aparicion >= 0:
        respuesta.append(aparicion)
        posicion_actual = aparicion + long
        aparicion = parrafo.find(palabra,posicion_actual)
    return respuesta
    
@login_required
def buscador_inteligente(request,tipo,investigacion_id,camino):
    """Genera la vista para las búsquedas inteligentes"""

    investigacion = Investigacion.objects.get(id=investigacion_id)
    if request.method == "POST":
        ids_eliminar = request.POST.getlist('checks2[]')
        for id in ids_eliminar:
             entidades = EntidadesDoc.objects.filter(id=id)
             for entidad in entidades:
                entidad.eliminado = True
                entidad.save()
    
    context = {"tipo":tipo,
                "investigacion_nombre":investigacion.nombre,
                "modelo":TipoModelo().getModeloString(investigacion.modelo),
                "org":"Organizaciones",
                "loc":"Locaciones",
                "lex":"Léxicos detectados",
                "pers":"Personas",
                "drog":"Drogas",
                "econ":"Económicos",
                "investigacion_id":investigacion.id,
                'title':'Buscador inteligente',
                }
    documentos = investigacion.documento_set.all().order_by('-fecha_agregado')
    resultado = []
    if tipo == "Léxicos detectados":
        for documento in documentos:
            tokens = TokensDoc.objects.filter(doc=documento).filter(eliminado=False)
            for t in tokens:
                parrafo = Parrafo.objects.get(id=t.parrafo.id)
                resultado.append([t.aparicion,t.categoria,t.lema,t.tipo,t.frase,parrafo.nro,parrafo.parrafo,documento.id,documento.nombre_doc,t.id])
    else:
        for documento in documentos:
            entidad = EntidadesDoc.objects.filter(doc=documento).filter(tipo=GetEntidad().getEntidad(tipo)).filter(eliminado=False)
            for e in entidad:
                parrafo = Parrafo.objects.get(id=e.parrafo.id)
                resultado.append([e.string, e.start, e.end, e.doc.id, documento.nombre_doc, parrafo.nro, parrafo.parrafo, e.id, e.string_original])
    context["res"] = resultado
    resultados_json = json.dumps(resultado)
    context["json"] = resultados_json
    breadcumb_path = parser_camino(camino)
    context["camino"] = camino
    context["camino_array"] = breadcumb_path
    return render(request,'FrontEnd/buscador_inteligente.html', context)

@login_required
def buscador_guiado(request,tipo,id_regex,investigacion_id,camino):
    """Genera la vista para búsquedas guiadas"""

    investigacion = Investigacion.objects.get(id=investigacion_id)
    context = {   "tipo":tipo,
                  "investigacion_nombre":investigacion.nombre,
                  "investigacion_id":investigacion.id,
                  'title':'Buscador guiado',
                  }

    if request.user.is_superuser:
        form_agregar = RegexForm()
        context['form_agregar'] = form_agregar

    if request.method == "POST":
        form = RegexForm(request.POST)
        if form.is_valid():
            nueva_regex = form.save(commit=False)
            nueva_regex.eliminado = False
            nueva_regex.save()
            return HttpResponseRedirect(reverse('buscador_guiado',args=[tipo,id_regex,investigacion_id,camino]))
    else:
        expresiones = Regex.objects.filter(eliminado=False).order_by('orden')
        context["expresiones"] = expresiones
        if tipo != "Búsqueda":
            print("entro")
            resultados = buscar_regex(tipo, investigacion, id_regex)
            d_res = defaultdict(list)
            for resultado in resultados:
                d_res[resultado[2]].append((resultado[0],resultado[1],resultado[3],resultado[4],resultado[5],resultado[6]))
            #resultado[0]: documento.id; resultado[1]: documento.nombre_doc; resultado[2]: m.group; resultado[3]: m.start(); resultado[4]: m.end(); 
            #resultad[5]: parrafo.nro; resultado[6]: parrafo.parrafo
            resultados = dict(d_res)
            context["res"] = resultados
            resultados_json = json.dumps(resultados)
            context["json"] = resultados_json
    breadcumb_path = parser_camino(camino)
    context["camino"] = camino
    context["camino_array"] = breadcumb_path
    return render(request,'FrontEnd/buscador_guiado.html', context)

def eliminar_regex(request,tipo,id_regex,investigacion_id,camino):
    """Elimina expresiones regulares"""
    ids_eliminar = request.POST.getlist('checks_regex[]')
    for id in ids_eliminar:
        regex = Regex.objects.get(id=id)
        regex.eliminado = True
        regex.save()
    return HttpResponseRedirect(reverse('buscador_guiado',args=[tipo,id_regex,investigacion_id,camino]))


def buscar_regex(tipo, investigacion, id_regex):
    """Busca dentro de un documento las expresiones regulares especificadas por parámetro"""

    #patron_url = "[-a-zA-Z0-9@:%._\+~#=]{2,256}\.[a-z]{2,6}\b([-a-zA-Z0-9@:%_\+.~#?&//=]*)"
    #patron_mail = "[^@\s]+@[^@\s]+\.[^@\s]+"
    #patron_tarjeta = "((\d{4} \d{6} \d{5})|(\d{4} \d{4} \d{4} \d{4}))"
    #patron_documento = "[1-9]{2,3}\.?[0-9]{3}\.?[0-9]{3}"
    #patron_telefono = "\+?([0-9]{1,2})?[ -]?9?[ -]?[0-9 \-]?[0-9]{6,13}"

    regex = Regex.objects.get(id=id_regex)
    patron_aux = fr"{regex.patron}"
    reg = re.compile(patron_aux)
    documentos = investigacion.documento_set.all().order_by('-fecha_agregado')
    resultados = []
    for documento in documentos:
        parrafos = Parrafo.objects.filter(doc_id= documento.id)
        for parrafo in parrafos:
            for m in reg.finditer(parrafo.parrafo):
                resultados.append((documento.id, documento.nombre_doc, m.group(), m.start(), m.end(), parrafo.nro, parrafo.parrafo))
    return resultados

def guardar_resultadoGeneral(request,investigacion_id,expresion,camino):
    """Guarda resultados de búsquedaa general"""
    if request.method == "POST":
        form = BuscadorGeneralForm()
        resultado_json = request.POST.get("resultado_json")
        resultado = json.loads(resultado_json)
        destacados = request.POST.getlist('checks3[]')
        investigacion = Investigacion.objects.get(id=investigacion_id)
        documentos = investigacion.documento_set.filter(eliminado=False)
        context = {'investigacion':investigacion,
                   'camino':camino,
                   'expresion':expresion,
                   'res':resultado,
                   'json':resultado_json,
                   'form':form,
                   'title':'Buscador general',
                   }
        resultado_header = ResultadoHeader(investigacion=investigacion,busqueda='General',tipo=expresion,estado=True,eliminado=False,propietario=request.user)
        resultado_header.save()
        for d in documentos:
            resultado_header.documentos.add(d)
        resultado_header.save()
        for fila in resultado:
            for aparicion in fila[4]:
                id = f'{fila[5]}.{aparicion}'
                documento = Documento.objects.get(id=fila[1])
                resultadoGen = ResultadoBusqGeneral(parrafo_nro=fila[2],posicion=aparicion,documento=documento,documento_nombre=fila[0],parrafo=fila[3],header=resultado_header)
                if id in destacados:
                    resultadoGen.destacado =True
                else:
                    resultadoGen.destacado = False
                resultadoGen.save()
        return render(request,'FrontEnd/buscador_general.html',context)
        
    return HttpResponseRedirect(reverse('buscador_general',args=[investigacion_id,camino]))

def guardar_resultadoGuiado(request,tipo,investigacion_id,camino):
    """Guarda resultados de búsqueda guiada"""
    if request.method == "POST":
        resultado_json = request.POST.get("resultado_json")
        resultado = json.loads(resultado_json)
        destacados = request.POST.getlist('checks3[]')
        investigacion = Investigacion.objects.get(id=investigacion_id)
        documentos = investigacion.documento_set.filter(eliminado=False)
        resultado_header = ResultadoHeader(investigacion=investigacion,busqueda='Guiada',tipo=tipo,estado=True,eliminado=False,propietario=request.user)
        resultado_header.save()
        for d in documentos:
            resultado_header.documentos.add(d)
        resultado_header.save()
        context = {"tipo":tipo,
                   "investigacion_nombre":investigacion.nombre,
                   "camino":camino,
                   "email":"Email",
                   "dni":"Documento",
                   "tarjeta":"Tarjeta",
                   "telefono":"Teléfono",
                   "url": "URL",
                   "investigacion_id":investigacion.id,
                   "res": resultado,
                   "title":'Buscador guiado',
                   "json":resultado_json,
                   }
        for key, value in resultado.items():
            header_general = ResultadoBusqGuiadaGeneral(clave=key,cantidadTotal=len(value),header=resultado_header)
            if key in destacados:
                header_general.destacado = True
            else:
                header_general.destacado = False
            header_general.save()
            for item in value:
                documento = Documento.objects.get(id=item[0])
                resultadoGui = ResultadoBusqGuiada(documento=documento,documento_nombre=item[1],start=item[2],end=item[3],parrafo_nro=item[4],parrafo=item[5],general=header_general)
                resultadoGui.save()
        return render(request,'FrontEnd/buscador_guiado.html', context)

        
    return HttpResponseRedirect(reverse('buscador_guiado',args=[tipo,investigacion_id,camino]))


def guardar_resultadoInteligente(request,tipo,investigacion_id,camino):
    """Guarda resultados de búsqueda inteligente"""
    if request.method == "POST":
        resultado_json = request.POST.get("resultado_json")
        resultado = json.loads(resultado_json)
        destacados = request.POST.getlist('checks3[]')
        investigacion = Investigacion.objects.get(id=investigacion_id)
        documentos = investigacion.documento_set.filter(eliminado=False)
        context = {"tipo":tipo,
                   "investigacion_nombre":investigacion.nombre,
                   "camino":camino,
                   "modelo":investigacion.modelo,
                   "org":"Organizaciones",
                   "loc":"Locaciones",
                   "lex":"Léxicos detectados",
                   "pers":"Personas",
                   "drog":"Drogas",
                   "econ":"Económicos",
                   "title":'Buscador inteligente',
                   "investigacion_id":investigacion.id,
                   "res":resultado,
                   "json":resultado_json,
                   }
        resultado_header = ResultadoHeader(investigacion=investigacion,busqueda='Inteligente',tipo=tipo,estado=True,eliminado=False,propietario=request.user)
        resultado_header.save()
        for d in documentos:
            resultado_header.documentos.add(d)
        resultado_header.save()
        for item in resultado:
            if tipo != "Léxicos detectados":
                entidad = EntidadesDoc.objects.get(id=item[7])
                resultadoInt = ResultadoBusqInteligente(doc=entidad.doc,string=entidad.string,string_original=entidad.string_original,start=entidad.start,end=entidad.end,parrafo_nro=entidad.parrafo.nro,parrafo=entidad.parrafo.parrafo,header=resultado_header)
                id = item[7] 
            else:
                token = TokensDoc.objects.get(id=item[9])
                resultadoInt = ResultadoBusqInteligenteTokens(doc=token.doc,aparicion=token.aparicion,tipo=token.tipo,frase=token.frase,lema=token.lema,categoria=token.categoria,parrafo=token.parrafo,parrafo_nro=token.parrafo.nro,header=resultado_header)
                id = item[9]
            if str(id) in destacados:
                resultadoInt.destacado = True
            else:
                resultadoInt.destacado = False
            resultadoInt.save()
        return render(request,'FrontEnd/buscador_inteligente.html', context)


    return HttpResponseRedirect(reverse('buscador_inteligente',args=[tipo,investigacion_id,camino]))

@login_required
def editar_entidad(request,tipo,id_ent,id_investigacion,camino):
    """Editar entidad"""
    if request.method == 'POST':
        string = request.POST.get('string')
        entidad = EntidadesDoc.objects.get(id=id_ent)
        entidad.string = string
        entidad.save()
    return HttpResponseRedirect(reverse('buscador_inteligente',args=[tipo,id_investigacion,camino]))


#*********************************************************************************************************************************************************************
#**************************************************************************************END BÚSQUEDAS *****************************************************************
#*********************************************************************************************************************************************************************

#*********************************************************************************************************************************************************************
#************************************************************************************** START RESULTADOS *************************************************************
#*********************************************************************************************************************************************************************

def generar_resultados(request,id_investigacion):
    """Genera los resultados para mostrar en las vistas"""
    investigacion = Investigacion.objects.get(id=id_investigacion)
    resultados = investigacion.resultadoheader_set.filter(eliminado=False).filter(propietario=request.user).order_by('-fecha')
    flag = False
    for resultado in resultados:
        if resultado.estado:
            documentos = investigacion.documento_set.filter(eliminado=False).values_list('id', flat=True)
            documentos_res = resultado.documentos.all().values_list('id',flat=True)
            if set(documentos) != set(documentos_res):
                flag = True
                resultado.estado = False
                resultado.save()
    if flag:
        resultados = investigacion.resultadoheader_set.filter(eliminado=False).order_by('-fecha')

    return resultados

@login_required
def resultados(request):
    """Resultados de las búsquedas"""
    form_elegir = BuscadorInvestigacionesForm(request.user)
    context = {
                'form_elegir':form_elegir,
                'title':'Resultados',
               }
    if request.method == 'POST':
        investigacion_id = request.POST.get('investigaciones')
        investigacion = Investigacion.objects.get(id=investigacion_id)
        camino = f'Resultados:resultados>{investigacion.nombre}:investigacion|{investigacion_id}'
        resultados = generar_resultados(request,investigacion_id)
        context['resultados'] = resultados
        context['investigacion_id'] = investigacion_id
        context['inicial'] = False
        context['nombre_investigacion'] = investigacion.nombre
        context['camino'] = camino

    return render(request,'FrontEnd/resultados.html', context)

@login_required
def resultados_investigacion(request,investigacion_id,destino,camino):
    """Resultados de una única investigación específica"""
    context = {
                'title':'Resultados investigacion',
               }
    investigacion = Investigacion.objects.get(id=investigacion_id)
    resultados = generar_resultados(request,investigacion_id)
    breadcrumb_path = parser_camino(camino)
    context['resultados'] = resultados
    context['investigacion_id'] = investigacion_id
    context['nombre_investigacion'] = investigacion.nombre
    context['camino'] = camino
    context['camino_array'] = breadcrumb_path

    return render(request,destino, context)

def eliminar_general_resultado(resultado_id,tipo):
    """Elimina resultados"""

    resultado = ResultadoHeader.objects.get(id=resultado_id)
    resultado.eliminado = True
    resultado.save()

    tipo_eliminar = {
            "Guiada":ResultadoBusqGuiadaGeneral,
            "General":ResultadoBusqGeneral,
            "Inteligente":ResultadoBusqInteligente,
        }

    #Eliminar físicamente todas las tuplas asociadas al resultado
    tuplas = tipo_eliminar[tipo].objects.filter(header=resultado)
    for tupla in tuplas:
        if tipo == "Guiada":
            guiadas = ResultadoBusqGuiada.objects.filter(general=tupla)
            guiadas.delete()
        tupla.delete()


@login_required
def eliminar_resultado(request,investigacion_id,resultado_id,tipo,camino):
    """Elimina resultado de búsqueda"""

    eliminar_general_resultado(resultado_id,tipo)

    resultados = generar_resultados(request,investigacion_id)
    form_elegir = BuscadorInvestigacionesForm(request.user)
    context = { 
                'title':'Resultados',
                'form_elegir':form_elegir,
                'resultados':resultados,
                'inicial':False,
                'investigacion_id': investigacion_id,
                'camino':camino,
               }

    return render(request,'FrontEnd/resultados.html', context)

@login_required
def eliminar_resultadoInvestigacion(request,investigacion_id,resultado_id,tipo,camino):
    """Elimina resultado de búsqueda"""

    eliminar_general_resultado(resultado_id,tipo)
    resultados = generar_resultados(request,investigacion_id)
    breadcrumb_path = parser_camino(camino)

    context = {
                'title':'Resultados',
                'resultados':resultados,
                'investigacion_id':investigacion_id,
                'camino':camino,
                'camino_array':breadcrumb_path,
               }

    return render(request,'FrontEnd/resultados_investigacion.html', context)

@login_required
def ver_resultado(request,resultado_id,tipo,camino):
    """Ver resultados específicos asociados a una búsqueda"""

    breadcrumb_path = parser_camino(camino)
    lexicos = "Léxicos detectados"
    resultado = ResultadoHeader.objects.get(id=resultado_id)
    context = {
                'header':resultado,
                'camino':camino,
                'camino_array':breadcrumb_path,
            }
    tipo_ver = {
            "Guiada":ResultadoBusqGuiadaGeneral,
            "General":ResultadoBusqGeneral,
            "Inteligente":ResultadoBusqInteligente,
        }

    if tipo == "Guiada":
        tuplas = tipo_ver[tipo].objects.filter(header=resultado)
        aux = []
        for tupla in tuplas:
            guiadas = ResultadoBusqGuiada.objects.filter(general=tupla)
            aux_guiadas = []
            for guiada in guiadas:
                aux_guiadas.append([guiada.documento.id,guiada.documento_nombre,guiada.start,guiada.end,guiada.parrafo,guiada.parrafo_nro])
            aux.append([tupla.clave,tupla.destacado,tupla.cantidadTotal,aux_guiadas])

        context['tuplas'] = aux
        context['title'] = 'Resultado guiado'
        return render(request,'FrontEnd/resultados_guiado.html', context)
    else:
        if resultado.tipo == lexicos:
            tuplas = ResultadoBusqInteligenteTokens.objects.filter(header=resultado)
            next = 'FrontEnd/resultados_inteligente_tokens.html'
        else:
            tuplas = tipo_ver[tipo].objects.filter(header=resultado)
            next = 'FrontEnd/resultados_inteligente.html'
        context['tuplas'] = tuplas

        if tipo == "Inteligente":
            context['title'] = 'Resultado inteligente'
            return render(request,next, context)
        else:
            context['title'] = 'Resultado general'
            return render(request,'FrontEnd/resultados_general.html', context)

        
#*********************************************************************************************************************************************************************
#************************************************************************************** END RESULTADOS ***************************************************************
#*********************************************************************************************************************************************************************


#*********************************************************************************************************************************************************************
#************************************************************************************** START INFORMES ***************************************************************
#*********************************************************************************************************************************************************************

@login_required
def crearInforme(request,resultado_id,tipo_informe,camino):
    """Crea informe en base a información de los resultados"""

    lexicos = "Léxicos detectados"

    resultado = ResultadoHeader.objects.get(id=resultado_id)
    busqueda = resultado.busqueda

    filepath = f'{os.path.dirname(os.path.dirname(os.path.abspath(__file__)))}\informes'

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'

    document.add_heading('Informe de búsqueda',0)

    busq = f'Búsqueda {busqueda}'
    p_usuario = document.add_paragraph('Usuario: ')
    p_usuario.add_run(request.user.username).italic = True
    paragraph_format = p_usuario.paragraph_format
    paragraph_format.space_before = Pt(50)
    p = document.add_paragraph('Informe generado en base a resultados obtenidos mediante una ')
    p.add_run(busq).italic = True
    p.add_run(' sobre la investigacion ')
    p.add_run(resultado.investigacion.nombre).italic = True
    p.add_run('.')
    p.style.font.size = Pt(13)
    
    p2 = document.add_paragraph('Fecha de búsqueda: ')
    p2.add_run(str(resultado.fecha.strftime("%d-%m-%Y %H:%M:%S")))

    p3 = document.add_paragraph('Las apariciones en ')
    p3.add_run('negrita').bold = True
    p3.add_run(' son aquellas que han sido señaladas como destacadas.')
    paragraph_format2 = p3.paragraph_format
    paragraph_format2.space_after = Pt(50)

    document.add_paragraph(f'"{resultado.tipo}"')

    tipo_busq = {
            "Guiada":ResultadoBusqGuiadaGeneral,
            "General":ResultadoBusqGeneral,
            "Inteligente":ResultadoBusqInteligente,
        }

    if tipo_informe == "destacados":
        if resultado.tipo != lexicos:
            tuplas = tipo_busq[busqueda].objects.filter(header=resultado).filter(destacado=True)
        else:
            tuplas = ResultadoBusqInteligenteTokens.objects.filter(header=resultado).filter(destacado=True)
    else:
        if resultado.tipo != lexicos:
            tuplas = tipo_busq[busqueda].objects.filter(header=resultado)
        else:
            tuplas = ResultadoBusqInteligenteTokens.objects.filter(header=resultado)

    if busqueda == "Guiada":
        for tupla in tuplas:
            guiadas = ResultadoBusqGuiada.objects.filter(general=tupla)
            if tupla.destacado:
                p4 = document.add_paragraph()
                p4.add_run(tupla.clave).bold = True
            else:
                p4 = document.add_paragraph(tupla.clave)
            paragraph_format = p4.paragraph_format
            paragraph_format.space_before = Pt(50)
            t = document.add_paragraph('Cantidad de apariciones: ')
            t.add_run(str(tupla.cantidadTotal))
            table = document.add_table(1,4)
            table.style = 'TableGrid'
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Documento'
            heading_cells[1].text = 'Nro. párrafo'
            heading_cells[2].text = 'Pos. inicial'
            heading_cells[3].text = 'Pos. final'
            for guiada in guiadas:
                cells = table.add_row().cells
                cells[0].text = guiada.documento_nombre
                cells[1].text = str(guiada.parrafo_nro)
                cells[2].text = str(guiada.start)
                cells[3].text = str(guiada.end)
    else:
        
        if busqueda == "Inteligente":
            if resultado.tipo != lexicos:
                table = document.add_table(1, 5)
                table.style = 'TableGrid'
                table.autofit = True
                heading_cells = table.rows[0].cells
                heading_cells[0].text = 'Aparición'
                heading_cells[1].text = 'Nro. párrafo'
                heading_cells[2].text = 'Pos. inicial'
                heading_cells[3].text = 'Pos. final'
                heading_cells[4].text = 'Documento'
                for tupla in tuplas:
                    cells = table.add_row().cells
                    if tupla.destacado:
                        cells[0].paragraphs[0].add_run(tupla.string).bold=True
                    else:
                        cells[0].text = tupla.string
                    cells[1].text = str(tupla.parrafo_nro)
                    cells[2].text = str(tupla.start)
                    cells[3].text = str(tupla.end)
                    cells[4].text = str(tupla.doc.nombre_doc)
            else:
                table = document.add_table(1, 6)
                table.style = 'TableGrid'
                table.autofit = True
                heading_cells = table.rows[0].cells
                heading_cells[0].text = 'Aparición'
                heading_cells[1].text = 'Base léxica'
                heading_cells[2].text = 'Categoría'
                heading_cells[3].text = 'Tipo'
                heading_cells[4].text = 'Documento'
                heading_cells[5].text = 'Nro. párrafo'
                for tupla in tuplas:
                    cells = table.add_row().cells
                    if tupla.destacado:
                        cells[0].paragraphs[0].add_run(tupla.aparicion).bold=True
                    else:
                        cells[0].text = tupla.aparicion
                    cells[1].text = str(tupla.lema)
                    cells[2].text = str(tupla.categoria)
                    cells[3].text = str(tupla.tipo)
                    cells[4].text = str(tupla.doc.nombre_doc)
                    cells[5].text = str(tupla.parrafo_nro)

        else:
            table = document.add_table(1, 4)
            table.style = 'TableGrid'
            table.autofit = True
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Aparición'
            heading_cells[1].text = 'Nro. párrafo'
            heading_cells[2].text = 'Pos. inicial'
            heading_cells[3].text = 'Documento'
            for tupla in tuplas:
                cells = table.add_row().cells
                if tupla.destacado:
                    cells[0].paragraphs[0].add_run(resultado.tipo).bold=True
                else:
                    cells[0].text = resultado.tipo
                cells[1].text = str(tupla.parrafo_nro)
                cells[2].text = str(tupla.posicion)
                cells[3].text = str(tupla.documento_nombre)
    
    informe = Informe(investigacion=resultado.investigacion,eliminado=False,busqueda=resultado.busqueda,propietario=request.user)
    informe.save()

    fecha = informe.fecha.strftime("%d-%m-%Y-%H-%M-%S")
    id = informe.id

    dominio = f'{filepath}\{id}-{fecha}.docx'

    informe.path = dominio
    informe.save()

    document.save(dominio)

    return HttpResponseRedirect(reverse('ver_resultado',args=[resultado_id,busqueda,camino]))

@login_required
def informes(request):
    """Muestra todos los informes para poder ser descargados"""
    form = BuscadorInvestigacionesForm(request.user)
    context = {'form':form, 
               'title':'Informes'
               }
    if request.method == 'POST':
        id_investigacion = request.POST.get('investigaciones')
        investigacion = Investigacion.objects.get(id=id_investigacion)
        informes = Informe.objects.filter(investigacion=investigacion).filter(eliminado=False).filter(propietario=request.user).order_by('-fecha')
        context['id_investigacion'] = id_investigacion
        context['informes'] = informes
        context['inicial'] = False
        context['nombre_investigacion'] = investigacion.nombre
    return render(request,'FrontEnd/informes.html',context)

@login_required
def informes_investigacion(request,investigacion_id,destino,camino):
    """Muestra todos los informes, de una investigación determinada, para poder ser descargados"""
    breadcrumb_path = parser_camino(camino)
    context = { 
               'title':'Informes',
               'camino':camino,
               'camino_array':breadcrumb_path,
               }
    investigacion = Investigacion.objects.get(id=investigacion_id)
    informes = Informe.objects.filter(investigacion=investigacion).filter(eliminado=False).filter(propietario=request.user).order_by('-fecha')
    context['id_investigacion'] = investigacion_id
    context['informes'] = informes
    return render(request,destino,context)

@login_required
def ver_informe(request, id_informe):
    """Descarga el informe para que el usuario pueda leerlo desde su dispositivo"""
    informe = Informe.objects.get(id=id_informe)
    nombre = f'Informe-{informe.fecha.strftime("%d-%m-%Y %H:%M:%S")}.docx'
    with open(informe.path, 'rb') as file:
        response = HttpResponse(file.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = 'attachment; filename=%s' % nombre
        return response
    raise Http404

@login_required
def eliminar_informe(request,informe_id):
    """Elimina informe"""
    
    if request.method == 'POST':
        next = request.POST.get('next', '/')

    informe = Informe.objects.get(id=informe_id)
    informe.eliminado = True
    informe.save()
    investigacion = informe.investigacion
    informes = Informe.objects.filter(investigacion=informe.investigacion).filter(eliminado=False).order_by('-fecha')
    form = BuscadorInvestigacionesForm(request.user)
    context = {'form':form, 
               'title':'Informes',
               'inicial':False,
               'id_investigacion':investigacion.id,
               'informes':informes,
               'nombre_investigacion':investigacion.nombre,
               }

    return render(request,next, context)


#*********************************************************************************************************************************************************************
#************************************************************************************** END INFORMES *****************************************************************
#*********************************************************************************************************************************************************************
